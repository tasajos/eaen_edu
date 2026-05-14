from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colores institucionales ──────────────────────────────────────────────────
NAVY     = RGBColor(0x0F, 0x17, 0x2A)   # fondo header
ORANGE   = RGBColor(0xEA, 0x58, 0x0C)   # acento naranja
GRAY_BG  = RGBColor(0xF1, 0xF5, 0xF9)   # fondo secciones
GRAY_TXT = RGBColor(0x64, 0x74, 0x8B)   # texto secundario
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BLACK    = RGBColor(0x1E, 0x29, 0x3B)


def set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_paragraph_border_left(para, color_hex="EA580C", size="18"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    size)
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def page_break(doc):
    doc.add_page_break()


def heading_section(doc, numero, titulo, subtitulo=""):
    """Encabezado de módulo con barra naranja izquierda."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    add_paragraph_border_left(p, "EA580C", "24")
    run_num = p.add_run(f"  {numero}.  ")
    run_num.bold      = True
    run_num.font.size = Pt(18)
    run_num.font.color.rgb = ORANGE
    run_tit = p.add_run(titulo)
    run_tit.bold      = True
    run_tit.font.size = Pt(18)
    run_tit.font.color.rgb = NAVY
    if subtitulo:
        ps = doc.add_paragraph(f"  {subtitulo}")
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after  = Pt(8)
        for run in ps.runs:
            run.italic = True
            run.font.color.rgb = GRAY_TXT
            run.font.size = Pt(11)


def sub_heading(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_TXT


def bullet(doc, texto, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def body(doc, texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK


def feature_table(doc, filas):
    """Tabla de 2 columnas: Funcionalidad | Descripción."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # Encabezado
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Funcionalidad", "Descripción"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], "0F172A")

    # Datos
    for idx, (func, desc) in enumerate(filas):
        row   = table.add_row().cells
        bg    = "F1F5F9" if idx % 2 == 0 else "FFFFFF"
        row[0].text = func
        row[1].text = desc
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(cell, bg)

    # Anchos
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("ESCUELA DE ALTOS ESTUDIOS NACIONALES")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("EAEN Avaroa")
run2.font.size = Pt(13)
run2.font.color.rgb = GRAY_TXT
run2.italic = True

doc.add_paragraph()

# Línea decorativa
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("─" * 50)
run3.font.color.rgb = ORANGE

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("SISTEMA DE GESTIÓN ACADÉMICA INSTITUCIONAL")
run4.bold = True
run4.font.size = Pt(20)
run4.font.color.rgb = NAVY

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Descripción Funcional de Módulos del Sistema")
run5.bold = True
run5.font.size = Pt(14)
run5.font.color.rgb = ORANGE

doc.add_paragraph()
doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run(f"Versión 1.0  ·  {datetime.date.today().strftime('%B %Y').capitalize()}")
run6.font.size = Pt(11)
run6.font.color.rgb = GRAY_TXT
run6.italic = True

page_break(doc)

# ── INTRODUCCIÓN ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run("Introducción")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

body(doc,
    "El Sistema de Gestión Académica Institucional de la Escuela de Altos Estudios Nacionales "
    "(EAEN Avaroa) es una plataforma web integral diseñada para centralizar y automatizar los "
    "procesos administrativos, académicos y financieros de la institución. El sistema permite la "
    "gestión de usuarios, promociones, calificaciones, asistencia, finanzas, evaluaciones y "
    "comunicaciones internas desde un entorno seguro y de acceso controlado por roles."
)
body(doc,
    "El presente documento describe de manera formal las funcionalidades de cada módulo disponible "
    "en el panel del Jefe de Estudios, así como el módulo de Perfil de Usuario accesible para todos "
    "los usuarios del sistema."
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 1 — Administración de Usuarios
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 1, "Administración de Usuarios",
    "Gestión integral del padrón de usuarios institucionales")

body(doc,
    "Este módulo centraliza el registro, consulta y actualización de todos los usuarios del sistema. "
    "Permite administrar el padrón de personal militar, civil y cursantes de postgrado, garantizando "
    "la integridad y trazabilidad de los datos institucionales."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Registro de usuarios",       "Alta de nuevos usuarios con datos de identificación, contacto, filiación militar y datos académicos."),
    ("Búsqueda por CI",            "Localización inmediata de cualquier usuario mediante su Carnet de Identidad."),
    ("Edición de datos",           "Actualización de todos los campos del usuario: nombre, apellidos, grado/profesión, filial, fuerza, turno, correo, teléfono, lugar de trabajo, etc."),
    ("Tipos de usuario",           "Gestión diferenciada por tipo: Cursante, Docente, Jefe de Curso, Jefe de Unidad, Finanzas y Administrador."),
    ("Control de estado",          "Activación o desactivación de cuentas. Los usuarios inactivos no pueden acceder al sistema."),
    ("Cambio de contraseña",       "Restablecimiento administrativo de contraseñas con encriptación segura (bcrypt)."),
    ("Listados filtrados",         "Consulta de cursantes activos, docentes y jefes de curso para asignación en otros módulos."),
])

sub_heading(doc, "Datos gestionados por usuario")
for campo in [
    "Apellido paterno, apellido materno y nombre",
    "Carnet de Identidad (CI) y lugar de expedición (EX)",
    "Grado militar o profesión civil",
    "Filial, fuerza y turno",
    "Fecha de nacimiento y fecha de inscripción",
    "Correo electrónico y email alternativo",
    "Teléfono y lugar de trabajo",
    "Rol y estado de cuenta",
]:
    bullet(doc, campo)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 2 — Administración de Promociones
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 2, "Administración de Promociones",
    "Creación y gestión de cursos y promociones académicas")

body(doc,
    "Este módulo permite la creación y administración completa de las promociones (cursos) "
    "ofertadas por la institución. Cada promoción agrupa a sus participantes, materias y docentes, "
    "configurando el entorno académico sobre el cual operan los demás módulos del sistema."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Creación de promociones",       "Registro de nuevos cursos con nombre, modalidad, horas académicas, fechas de inicio y fin, y estado (ACTIVO / INACTIVO)."),
    ("Gestión de participantes",      "Inscripción y baja de cursantes en una promoción. Control de participantes activos por curso."),
    ("Gestión de materias",           "Alta, edición y eliminación de materias vinculadas a una promoción, incluyendo código, descripción y carga horaria."),
    ("Asignación de docentes",        "Vinculación de un docente a cada materia dentro del curso."),
    ("Responsabilidades de curso",    "Asignación de roles adicionales dentro de una promoción (jefe de curso, secretario, etc.)."),
    ("Horarios",                      "Registro y gestión del horario de clases por materia dentro de cada promoción."),
    ("Vista del Jefe de Curso",       "Panel específico donde el Jefe de Curso puede consultar participantes, materias y estado de su promoción asignada."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 3 — Seguimiento Académico
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 3, "Seguimiento Académico",
    "Supervisión de calificaciones, asistencia y progreso académico")

body(doc,
    "El módulo de Seguimiento Académico centraliza todas las actividades relacionadas con la "
    "evaluación del desempeño académico de los cursantes. Permite a los docentes registrar "
    "calificaciones, controlar asistencia y gestionar tareas, mientras el Jefe de Estudios "
    "supervisa el progreso global de cada promoción."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Libro de calificaciones",     "Registro de notas parciales por materia con estructura de evaluación configurable (porcentajes por actividad)."),
    ("Notas finales",               "Cálculo automático de la nota final considerando calificaciones académicas y el componente disciplinario (méritos/deméritos)."),
    ("Control de asistencia",       "Registro de asistencia diaria por materia y cursante. Generación de resumen de presencia/ausencia."),
    ("Planificación de contenidos", "Registro del avance del sílabo por materia: temas planificados, ejecutados y pendientes."),
    ("Gestión de tareas",           "Creación de tareas por materia, recepción de entregas de cursantes (con archivo adjunto) y calificación de las mismas."),
    ("Evaluación de facilitadores", "Registro de la evaluación realizada por cursantes hacia los docentes, con resultado consolidado por materia."),
    ("Vista del cursante",          "Panel personal donde cada cursante consulta sus calificaciones, asistencias, tareas pendientes y estado financiero."),
    ("Vista del docente",           "Acceso del docente a su libro de calificaciones, asistencia, tareas y planificación por materia asignada."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 4 — Disciplina
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 4, "Disciplina",
    "Registro de méritos y deméritos con impacto en la nota final")

body(doc,
    "El módulo de Disciplina gestiona el componente disciplinario del desempeño académico. "
    "Permite registrar méritos y deméritos a los cursantes, los cuales inciden directamente "
    "en la nota final de cada materia conforme al porcentaje institucional establecido."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Catálogo de conductas",       "Listado oficial de acciones que generan mérito o demérito, con su valor en puntos correspondiente."),
    ("Registro de incidencias",     "Alta de méritos o deméritos a un cursante específico dentro de una materia y curso determinados."),
    ("Historial disciplinario",     "Consulta del historial completo de méritos y deméritos de un cursante por curso, con posibilidad de eliminación de registros."),
    ("Resumen por curso",           "Vista consolidada del comportamiento disciplinario de todos los participantes de una promoción."),
    ("Impacto en nota final",       "El puntaje disciplinario acumulado se integra automáticamente al cálculo de la nota final de cada materia."),
    ("Acceso por rol",              "El Jefe de Curso puede registrar y consultar incidencias desde su propio panel de administración."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 5 — Administración Finanzas
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 5, "Administración Finanzas",
    "Control de pagos, conceptos y bloqueo automático por mora")

body(doc,
    "El módulo de Administración Finanzas permite gestionar los conceptos de cobro institucional "
    "y llevar el registro de pagos por usuario. Cuenta con un mecanismo automático de bloqueo "
    "de acceso para cursantes con pagos pendientes, garantizando el cumplimiento de las "
    "obligaciones financieras como requisito de uso del sistema."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Gestión de conceptos",        "Creación de conceptos de cobro: Matrícula, Guía y Mensualidades, con montos, periodos y descripción."),
    ("Registro de pagos",           "Marcado de pagos como realizados por usuario y concepto. Registro de fecha y estado del pago."),
    ("Estado financiero",           "Consulta del estado de cuenta de cada cursante: conceptos pagados, pendientes y monto total adeudado."),
    ("Bloqueo automático",          "Los cursantes con pagos pendientes ven bloqueado su acceso al sistema con un aviso detallado de los montos adeudados."),
    ("Resumen financiero global",   "Vista consolidada de todos los usuarios con sus estados de cuenta para seguimiento por parte del área de finanzas."),
    ("Panel del área de finanzas",  "Acceso diferenciado para el personal de finanzas que permite gestionar pagos sin acceso a otros módulos académicos."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 6 — Evaluaciones Institucionales
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 6, "Evaluaciones Institucionales",
    "Gestión de evaluaciones de cursantes a docentes y entre pares")

body(doc,
    "El módulo de Evaluaciones Institucionales permite diseñar, habilitar y gestionar procesos "
    "de evaluación formal entre los actores del sistema. Principalmente se utiliza para la "
    "evaluación de docentes (facilitadores) por parte de los cursantes, generando resultados "
    "consolidados para el análisis institucional."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Plantillas de evaluación",    "Configuración de los cuestionarios de evaluación con sus indicadores, criterios y escala de puntuación."),
    ("Periodos de evaluación",      "Creación de periodos que habilitan o deshabilitan la ventana de tiempo en que los usuarios pueden responder evaluaciones."),
    ("Habilitación / cierre",       "Activación y desactivación de periodos de evaluación con un solo clic desde el panel administrativo."),
    ("Evaluación de facilitadores", "Formulario que cada cursante completa para calificar el desempeño de sus docentes en cada materia."),
    ("Evaluaciones pendientes",     "Notificación y acceso directo a las evaluaciones que el usuario aún no ha completado."),
    ("Resultados por periodo",      "Consulta de resultados consolidados por periodo: promedios por indicador, por docente y por curso."),
    ("Resultados de cursantes",     "Vista de las respuestas emitidas por los cursantes en un periodo determinado."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 7 — Administración de Notificaciones
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 7, "Administración de Notificaciones",
    "Comunicación institucional hacia todos los usuarios del sistema")

body(doc,
    "Este módulo permite al Jefe de Estudios emitir comunicados y alertas institucionales "
    "dirigidos a todos los usuarios del sistema. Las notificaciones se muestran en tiempo real "
    "en los paneles de cada usuario con indicador visual de mensajes no leídos."
)

sub_heading(doc, "Funcionalidades principales")
feature_table(doc, [
    ("Redacción de notificaciones",  "Creación de mensajes con título, contenido y fecha de emisión."),
    ("Difusión institucional",       "Las notificaciones son visibles para todos los usuarios con sesión activa en el sistema."),
    ("Indicador de no leídas",       "Cada usuario visualiza un contador de notificaciones pendientes de lectura en su panel principal."),
    ("Marcar como leída",            "El usuario puede marcar individualmente cada notificación como leída para limpiar su bandeja."),
    ("Marcar todas como leídas",     "Acción masiva para limpiar todas las notificaciones de una vez."),
    ("Eliminación de notificaciones","El administrador puede eliminar notificaciones obsoletas o enviadas por error."),
    ("Estadísticas",                 "Vista del total de notificaciones emitidas y su estado de lectura global."),
])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# MÓDULO 8 — Perfil de Usuario
# ════════════════════════════════════════════════════════════════════════════
heading_section(doc, 8, "Perfil de Usuario",
    "Consulta de datos personales y cambio de contraseña")

body(doc,
    "El módulo de Perfil de Usuario es accesible para todos los roles del sistema desde sus "
    "respectivos paneles. Permite al usuario consultar su información personal registrada en el "
    "sistema y gestionar su contraseña de acceso de forma autónoma y segura."
)

sub_heading(doc, "Pestaña: Mis Datos")
body(doc,
    "Muestra la información completa del usuario organizada en tres secciones:"
)
bullet(doc, "Identificación: CI, lugar de expedición (EX), apellidos, nombre, grado/profesión y fecha de nacimiento.")
bullet(doc, "Contacto y Trabajo: correo electrónico, email alternativo, teléfono, lugar de trabajo, filial, fuerza y turno.")
bullet(doc, "Inscripción: fecha de inscripción y estado de la cuenta.")

doc.add_paragraph()
sub_heading(doc, "Pestaña: Cambiar Contraseña")
body(doc,
    "Permite al usuario actualizar su contraseña de acceso de forma segura:"
)
feature_table(doc, [
    ("Verificación de identidad",   "El usuario debe ingresar su contraseña actual como confirmación antes de establecer una nueva."),
    ("Nueva contraseña",            "Ingreso de la nueva contraseña con un mínimo de 6 caracteres."),
    ("Confirmación",                "Campo de confirmación que debe coincidir exactamente con la nueva contraseña para evitar errores de escritura."),
    ("Visibilidad de contraseña",   "Botón de ojo en cada campo para mostrar u ocultar el texto ingresado."),
    ("Encriptación segura",         "La contraseña se almacena siempre de forma encriptada (bcrypt) en la base de datos. Nunca en texto plano."),
    ("Feedback inmediato",          "Mensajes de éxito o error mostrados al usuario tras el intento de actualización."),
])

page_break(doc)

# ── CIERRE ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run("Consideraciones Generales del Sistema")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

feature_table(doc, [
    ("Control de acceso por roles",  "Cada módulo es accesible únicamente por los roles autorizados. El sistema valida el rol en cada sesión."),
    ("Bloqueo financiero",           "Los cursantes con deuda vencida quedan bloqueados automáticamente, independientemente del módulo al que intenten acceder."),
    ("Sesión segura",                "La sesión del usuario se almacena localmente y expira al cerrar sesión o al limpiar el navegador."),
    ("API REST",                     "Todos los módulos se comunican con el servidor a través de una API REST centralizada, garantizando consistencia de datos."),
    ("Responsive",                   "La interfaz es adaptable a dispositivos móviles y de escritorio, con menú lateral colapsable en pantallas pequeñas."),
])

doc.add_paragraph()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_fin = p_fin.add_run(
    f"Escuela de Altos Estudios Nacionales — EAEN Avaroa\n"
    f"Documento generado el {datetime.date.today().strftime('%d de %B de %Y')}"
)
run_fin.italic = True
run_fin.font.size = Pt(10)
run_fin.font.color.rgb = GRAY_TXT

# ── GUARDAR ──────────────────────────────────────────────────────────────────
output = "EAEN_Descripcion_Funcional_Modulos.docx"
doc.save(output)
print(f"Documento generado: {output}")
