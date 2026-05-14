from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Paleta ───────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0F, 0x17, 0x2A)
ORANGE    = RGBColor(0xEA, 0x58, 0x0C)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TXT  = RGBColor(0x64, 0x74, 0x8B)
BLACK     = RGBColor(0x1E, 0x29, 0x3B)
GREEN_D   = RGBColor(0x16, 0x6A, 0x34)
GREEN_L   = RGBColor(0xDC, 0xFC, 0xE7)
RED_D     = RGBColor(0x99, 0x1B, 0x1B)
RED_L     = RGBColor(0xFE, 0xF2, 0xF2)

ROLE_COLORS = {
    "JEFE_ESTUDIOS": ("0F172A", "FFFFFF"),   # navy / white
    "DOCENTE":       ("1D4ED8", "FFFFFF"),   # blue / white
    "CURSANTE":      ("166534", "FFFFFF"),   # green / white
    "JEFE_CURSO":    ("92400E", "FFFFFF"),   # amber / white
    "ADMIN_FINANZAS":("7C3AED", "FFFFFF"),   # purple / white
}

# ── Helpers XML / estilos ─────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_left_border(para, color="EA580C", sz="20"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def body(doc, texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK


def bullet(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def sub_label(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_TXT


def modulo_table(doc, filas):
    """Tabla de módulos: Módulo | Acciones disponibles."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    hdr = table.rows[0].cells
    for i, txt in enumerate(["Módulo / Sección", "Acciones disponibles"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], "0F172A")

    for idx, (modulo, acciones) in enumerate(filas):
        row = table.add_row().cells
        bg  = "F1F5F9" if idx % 2 == 0 else "FFFFFF"
        row[0].text = modulo
        row[1].text = acciones
        for j, cell in enumerate(row):
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True
                run.font.color.rgb = NAVY
            set_cell_bg(cell, bg)

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)

    doc.add_paragraph()


def role_header(doc, rol_key, titulo, subtitulo, descripcion):
    """Bloque de cabecera de rol con color propio."""
    bg, fg = ROLE_COLORS.get(rol_key, ("0F172A", "FFFFFF"))

    # Tabla 1-celda como banner
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Cm(16)

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(0)
    r1 = p1.add_run(titulo)
    r1.bold = True
    r1.font.size = Pt(18)
    r1.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    r2 = p2.add_run(subtitulo)
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()
    body(doc, descripcion)


# ════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("ESCUELA DE ALTOS ESTUDIOS NACIONALES")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("EAEN Avaroa")
r2.italic = True; r2.font.size = Pt(13); r2.font.color.rgb = GRAY_TXT

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("─" * 50)
r3.font.color.rgb = ORANGE

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("SISTEMA DE GESTIÓN ACADÉMICA INSTITUCIONAL")
r4.bold = True; r4.font.size = Pt(20); r4.font.color.rgb = NAVY

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("Descripción de Roles y Accesos por Usuario")
r5.bold = True; r5.font.size = Pt(14); r5.font.color.rgb = ORANGE

doc.add_paragraph()
doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run(f"Versión 1.0  ·  {datetime.date.today().strftime('%B %Y').capitalize()}")
r6.italic = True; r6.font.size = Pt(11); r6.font.color.rgb = GRAY_TXT

doc.add_page_break()

# ── INTRODUCCIÓN ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("Introducción")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

body(doc,
    "El sistema de la EAEN Avaroa implementa un control de acceso basado en roles (RBAC). "
    "Cada usuario accede únicamente a los módulos y acciones que corresponden a su función "
    "institucional, garantizando la confidencialidad de la información y la integridad de los "
    "procesos académicos, financieros y disciplinarios."
)
body(doc,
    "Este documento describe, para cada rol del sistema, su perfil institucional, los módulos "
    "a los que tiene acceso y las acciones específicas que puede realizar dentro de cada uno."
)

sub_label(doc, "Roles del sistema")
for rol, desc in [
    ("Jefe de Estudios / Administrador", "Acceso total al sistema. Gestiona todos los módulos."),
    ("Docente",                          "Acceso a sus materias asignadas: asistencia, calificaciones, tareas y evaluaciones."),
    ("Cursante",                         "Acceso a su información académica, financiera y disciplinaria personal."),
    ("Jefe de Curso",                    "Supervisión de los participantes y materias de su promoción asignada."),
    ("Administrador de Finanzas",        "Acceso exclusivo al módulo de finanzas para gestión de pagos."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{rol}: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = BLACK

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ROL 1 — JEFE DE ESTUDIOS / ADMINISTRADOR
# ════════════════════════════════════════════════════════════════════════════
role_header(doc,
    "JEFE_ESTUDIOS",
    "Jefe de Estudios / Administrador",
    "Rol: JEFE_ESTUDIOS / ADMIN  ·  Acceso: Panel principal del sistema",
    "Es el rol con mayor nivel de privilegio en el sistema. Tiene acceso completo a todos los "
    "módulos de administración, configuración y supervisión. Desde su panel puede gestionar "
    "usuarios, promociones, finanzas, evaluaciones, disciplina y comunicaciones institucionales."
)

sub_label(doc, "Módulos y acciones disponibles")
modulo_table(doc, [
    ("Administración\nde Usuarios",
     "Registrar, buscar, editar y desactivar usuarios de todos los roles. "
     "Restablecer contraseñas. Consultar listados filtrados por tipo."),
    ("Administración\nde Promociones",
     "Crear y editar cursos/promociones. Inscribir y dar de baja participantes. "
     "Gestionar materias, asignar docentes, definir horarios y responsabilidades."),
    ("Seguimiento Académico",
     "Supervisar calificaciones, asistencia y avance académico de todas las materias y promociones. "
     "Consultar libro de notas, resúmenes de tareas y evaluaciones de facilitadores."),
    ("Disciplina",
     "Configurar el porcentaje disciplinario por materia. Consultar y gestionar "
     "el historial de méritos y deméritos de todos los cursantes."),
    ("Administración Finanzas",
     "Crear conceptos de cobro (matrícula, guía, mensualidades). Registrar y consultar pagos. "
     "Ver el estado financiero de todos los cursantes. Gestionar bloqueos por mora."),
    ("Evaluaciones\nInstitucionales",
     "Configurar plantillas de evaluación. Crear y habilitar/cerrar periodos de evaluación. "
     "Consultar resultados consolidados por periodo, materia y docente."),
    ("Administración\nde Notificaciones",
     "Redactar y publicar notificaciones institucionales visibles para todos los usuarios. "
     "Eliminar comunicados obsoletos. Consultar estadísticas de lectura."),
    ("Perfil de Usuario",
     "Consultar sus propios datos personales. Cambiar su contraseña de acceso."),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ROL 2 — DOCENTE
# ════════════════════════════════════════════════════════════════════════════
role_header(doc,
    "DOCENTE",
    "Docente",
    "Rol: DOCENTE  ·  Acceso: Panel del Docente",
    "El Docente accede a un panel personalizado que muestra exclusivamente las materias que "
    "le han sido asignadas dentro de cada promoción. Desde allí gestiona la asistencia, "
    "calificaciones, tareas y evaluación de su desempeño como facilitador."
)

sub_label(doc, "Módulos y acciones disponibles")
modulo_table(doc, [
    ("Asistencia",
     "Registrar la asistencia diaria de los cursantes en cada una de sus materias asignadas. "
     "Consultar resumen de presencia y ausencias por participante."),
    ("Calificaciones",
     "Registrar y actualizar notas parciales en el libro de calificaciones de su materia. "
     "Consultar el promedio académico y la nota final de cada cursante."),
    ("Tareas",
     "Crear tareas para su materia con fecha límite y descripción. "
     "Revisar entregas de los cursantes (con archivos adjuntos). "
     "Calificar cada entrega y registrar observaciones."),
    ("Facilitador",
     "Consultar los resultados de la evaluación institucional que los cursantes realizaron "
     "sobre su desempeño como facilitador de la materia."),
    ("Calendario",
     "Visualizar el calendario académico de su materia con las actividades planificadas."),
    ("Notificaciones",
     "Recibir y consultar las notificaciones institucionales emitidas por el Jefe de Estudios. "
     "Marcar notificaciones como leídas."),
    ("Perfil de Usuario",
     "Consultar sus datos personales registrados en el sistema. Cambiar su contraseña."),
])

sub_label(doc, "Restricciones")
bullet(doc, "Solo puede ver y gestionar las materias que le han sido asignadas por el administrador.")
bullet(doc, "No tiene acceso a información financiera, disciplinaria ni a datos de otros docentes.")
bullet(doc, "No puede crear ni modificar usuarios, cursos ni notificaciones institucionales.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ROL 3 — CURSANTE
# ════════════════════════════════════════════════════════════════════════════
role_header(doc,
    "CURSANTE",
    "Cursante",
    "Rol: CURSANTE  ·  Acceso: Panel del Cursante",
    "El Cursante accede a un panel personal de solo consulta donde puede revisar su situación "
    "académica, financiera y disciplinaria. También puede participar en evaluaciones "
    "institucionales, entregar tareas y consultar su calendario de actividades."
)

sub_label(doc, "Módulos y acciones disponibles")
modulo_table(doc, [
    ("Mis Notas / Evaluaciones",
     "Consultar sus calificaciones parciales y nota final por materia. "
     "Ver el resumen académico general de todas sus materias inscritas."),
    ("Tareas",
     "Ver las tareas asignadas por sus docentes en cada materia. "
     "Entregar tareas en formato de archivo adjunto antes del plazo límite. "
     "Consultar la calificación y observaciones recibidas por cada entrega."),
    ("Asistencia",
     "Consultar su registro de asistencia en cada materia: fechas presentes, ausentes y tardanzas."),
    ("Disciplina",
     "Consultar su historial de méritos y deméritos registrados dentro de su promoción, "
     "y el impacto acumulado en su nota final."),
    ("Pagos",
     "Revisar el estado de sus obligaciones financieras: conceptos pagados y pendientes. "
     "Si tiene deuda vencida, el acceso al sistema queda bloqueado automáticamente con "
     "detalle de los montos adeudados."),
    ("Evaluaciones\nInstitucionales",
     "Completar las evaluaciones institucionales habilitadas para valorar el desempeño "
     "de sus docentes (facilitadores) durante el periodo activo."),
    ("Calendario",
     "Visualizar el calendario académico de su curso con fechas relevantes."),
    ("Notificaciones",
     "Recibir y consultar notificaciones institucionales. Marcar como leídas."),
    ("Perfil de Usuario",
     "Consultar sus datos personales. Cambiar su contraseña de acceso."),
])

sub_label(doc, "Restricciones")
bullet(doc, "Solo puede ver su propia información: no accede a datos de otros cursantes.")
bullet(doc, "No puede modificar calificaciones, asistencia ni registros disciplinarios.")
bullet(doc, "El bloqueo financiero por mora impide el acceso a todos los módulos hasta regularizar el pago.")
bullet(doc, "No tiene acceso a módulos de administración, gestión de usuarios ni notificaciones institucionales.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ROL 4 — JEFE DE CURSO
# ════════════════════════════════════════════════════════════════════════════
role_header(doc,
    "JEFE_CURSO",
    "Jefe de Curso",
    "Rol: JEFE_CURSO  ·  Acceso: Panel del Jefe de Curso",
    "El Jefe de Curso es el responsable operativo de una promoción específica. Su panel le "
    "permite supervisar el estado general del curso: participantes, materias asignadas y "
    "situación disciplinaria. También puede gestionar las evaluaciones institucionales."
)

sub_label(doc, "Módulos y acciones disponibles")
modulo_table(doc, [
    ("Resumen del Curso",
     "Visualizar información general de su promoción: nombre, modalidad, horas académicas, "
     "fechas de inicio y fin, estado activo/inactivo y estadísticas de participantes y materias."),
    ("Participantes",
     "Consultar el listado completo de cursantes inscritos en su curso con datos de "
     "identificación (nombre, CI, correo) y estado de la cuenta."),
    ("Materias",
     "Ver las materias del curso con su descripción, carga horaria y el docente asignado a cada una. "
     "Identificar materias sin docente asignado."),
    ("Disciplina",
     "Registrar méritos y deméritos a los participantes de su curso. "
     "Consultar el historial disciplinario individual y el resumen por curso."),
    ("Evaluaciones\nInstitucionales",
     "Crear y gestionar periodos de evaluación institucional para su curso. "
     "Habilitar o cerrar evaluaciones. Consultar resultados consolidados."),
    ("Notificaciones",
     "Recibir y consultar notificaciones institucionales. Marcar como leídas."),
    ("Perfil de Usuario",
     "Consultar sus datos personales. Cambiar su contraseña de acceso."),
])

sub_label(doc, "Restricciones")
bullet(doc, "Solo tiene visibilidad sobre el curso o cursos que le han sido asignados como Jefe.")
bullet(doc, "No puede modificar datos de usuarios, calificaciones ni información financiera.")
bullet(doc, "No tiene acceso a los módulos de administración central del Jefe de Estudios.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ROL 5 — ADMIN FINANZAS
# ════════════════════════════════════════════════════════════════════════════
role_header(doc,
    "ADMIN_FINANZAS",
    "Administrador de Finanzas",
    "Rol: ADMIN_FINANZAS  ·  Acceso: Módulo de Finanzas",
    "El Administrador de Finanzas accede directamente al módulo de gestión financiera, sin "
    "visibilidad sobre otros módulos académicos o administrativos del sistema. Su función "
    "es el control y registro de los pagos institucionales."
)

sub_label(doc, "Módulos y acciones disponibles")
modulo_table(doc, [
    ("Conceptos de Cobro",
     "Crear y eliminar conceptos de cobro institucionales: Matrícula, Guía y Mensualidades, "
     "con montos, descripción y periodos correspondientes."),
    ("Registro de Pagos",
     "Marcar conceptos como pagados por usuario. Registrar la fecha de pago. "
     "Actualizar el estado de cada obligación financiera."),
    ("Estado Financiero\nde Cursantes",
     "Consultar el estado de cuenta individual de cada cursante: conceptos pagados, "
     "pendientes y monto total adeudado."),
    ("Resumen Global",
     "Vista consolidada de todos los usuarios con el estado de sus obligaciones financieras "
     "para seguimiento y control del área."),
    ("Bloqueo por Mora",
     "El sistema aplica automáticamente el bloqueo de acceso a cursantes con pagos vencidos. "
     "Al registrar el pago correspondiente, el acceso se restaura de forma inmediata."),
    ("Perfil de Usuario",
     "Consultar sus propios datos personales. Cambiar su contraseña de acceso."),
])

sub_label(doc, "Restricciones")
bullet(doc, "No tiene acceso a módulos académicos: calificaciones, asistencia, tareas ni disciplina.")
bullet(doc, "No puede gestionar usuarios, cursos, notificaciones ni evaluaciones institucionales.")
bullet(doc, "Su acceso es exclusivo al módulo de finanzas desde el momento del inicio de sesión.")

doc.add_page_break()

# ── TABLA RESUMEN COMPARATIVA ─────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("Cuadro Comparativo de Acceso por Rol")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(10)

MODULOS = [
    "Adm. Usuarios",
    "Adm. Promociones",
    "Seguimiento Académico",
    "Disciplina",
    "Adm. Finanzas",
    "Eval. Institucionales",
    "Notificaciones",
    "Perfil de Usuario",
]

ROLES_COLS = [
    "Jefe\nEstudios",
    "Docente",
    "Cursante",
    "Jefe\nCurso",
    "Adm.\nFinanzas",
]

# Acceso: ✔ total | ◑ parcial | ✘ sin acceso
MATRIZ = [
    # Jefe  Docente  Cursante  JefeCurso  Finanzas
    ["✔",   "✘",     "✘",      "✘",       "✘"],   # Adm. Usuarios
    ["✔",   "✘",     "✘",      "✘",       "✘"],   # Adm. Promociones
    ["✔",   "◑",     "◑",      "✘",       "✘"],   # Seguimiento Académico
    ["✔",   "✘",     "◑",      "✔",       "✘"],   # Disciplina
    ["✔",   "✘",     "◑",      "✘",       "✔"],   # Adm. Finanzas
    ["✔",   "◑",     "◑",      "✔",       "✘"],   # Eval. Institucionales
    ["✔",   "◑",     "◑",      "◑",       "✘"],   # Notificaciones
    ["✔",   "✔",     "✔",      "✔",       "✔"],   # Perfil de Usuario
]

cols = 1 + len(ROLES_COLS)
table = doc.add_table(rows=1 + len(MODULOS), cols=cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# Encabezados
hdr = table.rows[0].cells
hdr[0].text = "Módulo"
set_cell_bg(hdr[0], "0F172A")
hdr[0].paragraphs[0].runs[0].bold = True
hdr[0].paragraphs[0].runs[0].font.color.rgb = WHITE
hdr[0].paragraphs[0].runs[0].font.size = Pt(9)

ROLE_BG = ["0F172A", "1D4ED8", "166534", "92400E", "7C3AED"]
for j, (rol_txt, bg) in enumerate(zip(ROLES_COLS, ROLE_BG)):
    c = hdr[j + 1]
    c.text = rol_txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = WHITE
    c.paragraphs[0].runs[0].font.size = Pt(9)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, bg)

# Filas de datos
for i, (modulo, fila) in enumerate(zip(MODULOS, MATRIZ)):
    row = table.rows[i + 1].cells
    bg_row = "F8FAFC" if i % 2 == 0 else "FFFFFF"
    row[0].text = modulo
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[0].paragraphs[0].runs[0].bold = True
    row[0].paragraphs[0].runs[0].font.color.rgb = NAVY
    set_cell_bg(row[0], bg_row)
    for j, val in enumerate(fila):
        c = row[j + 1]
        c.text = val
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].runs[0]
        run.font.size = Pt(12)
        run.bold = True
        if val == "✔":
            run.font.color.rgb = GREEN_D
            set_cell_bg(c, "F0FDF4")
        elif val == "◑":
            run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            set_cell_bg(c, "FFFBEB")
        else:
            run.font.color.rgb = RED_D
            set_cell_bg(c, "FEF2F2")

# Anchos de columna
for row in table.rows:
    row.cells[0].width = Cm(4.8)
    for j in range(1, cols):
        row.cells[j].width = Cm(2.0)

doc.add_paragraph()

# Leyenda
p_ley = doc.add_paragraph()
p_ley.paragraph_format.space_before = Pt(4)
for sym, desc, color in [
    ("✔", " Acceso completo    ", GREEN_D),
    ("◑", " Acceso parcial / solo lectura    ", RGBColor(0x92, 0x40, 0x0E)),
    ("✘", " Sin acceso", RED_D),
]:
    r = p_ley.add_run(sym)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = color
    r2 = p_ley.add_run(desc)
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY_TXT

# ── PIE DE PÁGINA ─────────────────────────────────────────────────────────────
doc.add_paragraph()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = p_fin.add_run(
    f"Escuela de Altos Estudios Nacionales — EAEN Avaroa\n"
    f"Documento generado el {datetime.date.today().strftime('%d de %B de %Y')}"
)
r_fin.italic = True; r_fin.font.size = Pt(10); r_fin.font.color.rgb = GRAY_TXT

# ── GUARDAR ───────────────────────────────────────────────────────────────────
output = "EAEN_Roles_y_Accesos.docx"
doc.save(output)
print(f"Documento generado: {output}")
